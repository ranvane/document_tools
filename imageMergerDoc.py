#!/usr/bin/env python
# -*- coding: utf-8 -*-
import wx
import os
from docx import Document
from docx.shared import Cm
from PIL import Image
from loguru import logger

from imageMergerDoc_UI import Main_Ui_Frame  # 导入生成的界面类
from FileDropTarget import FileDropTarget  # 导入文件拖放类
class MainFrame(Main_Ui_Frame):
    """
        主应用程序类：将图片插入到 Word 文档中。

        功能：
        - 拖放或选择多个图片文件。
        - 显示图片列表并提供预览功能。
        - 将所有图片生成一个 Word (.docx) 文档，并保存到指定路径。

        依赖库：
        - wxPython: GUI 界面框架。
        - python-docx: 用于创建和编辑 Word 文档。
        - Pillow (PIL): 图片处理。
    """
    def __init__(self, parent):
        super().__init__(parent)
        # 加载图标
        try:
            icon_path = os.path.join(os.path.dirname(__file__), "imageMergerDoc_icon.ico")
            if os.path.exists(icon_path):
                icon = wx.Icon(icon_path, wx.BITMAP_TYPE_ICO)
                self.SetIcon(icon)
        except Exception as e:
            logger.error(f"Failed to load icon: {e}")
            
        self.image_paths = []  # 存储已选图片路径
        

        
    def on_select_files(self, event):
        """打开文件对话框，让用户选择多个图片文件"""
        with wx.FileDialog(self, "选择图片", wildcard="Image files (*.png;*.jpg;*.jpeg)|*.png;*.jpg;*.jpeg",
                           style=wx.FD_OPEN | wx.FD_MULTIPLE) as fileDialog:
            if fileDialog.ShowModal() == wx.ID_CANCEL:
                return
            paths = fileDialog.GetPaths()
            self.add_images(paths)

    def add_images(self, paths):
        """将有效图片路径添加到列表中，并更新 ListBox 显示"""
        for path in paths:
            if path.lower().endswith(('.png', '.jpg', '.jpeg')) and path not in self.image_paths:
                self.image_paths.append(path)
                self.m_ImageListBox.Append(path)

    def on_delete_selected(self, event):
        """从列表中删除选中的图片项"""
        selections = self.m_ImageListBox.GetSelections()
        if not selections:
            return
        for index in reversed(selections):
            del self.image_paths[index]
            self.m_ImageListBox.Delete(index)
        self.preview.SetBitmap(wx.NullBitmap)

    def on_delete_all(self, event):
        """清空所有图片列表"""
        self.image_paths.clear()
        self.m_ImageListBox.Clear()
        self.preview.SetBitmap(wx.NullBitmap)

    def on_preview_image(self, event):
        """在右侧显示所选图片的预览"""
        index = event.GetSelection()
        if 0 <= index < len(self.image_paths):
            path = self.image_paths[index]
            try:
                image = wx.Image(path, wx.BITMAP_TYPE_ANY)
                image = image.Scale(400, 400, wx.IMAGE_QUALITY_HIGH)
                self.preview.SetBitmap(wx.Bitmap(image))
            except Exception as e:
                wx.MessageBox(f"预览失败：{e}", "错误", wx.ICON_ERROR)

    def on_generate_doc(self, event):
        """将所有图片插入 Word 文档，并提示用户保存"""
        if not self.image_paths:
            wx.MessageBox("请先添加图片！", "提示", wx.ICON_INFORMATION)
            return

        doc = Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.bottom_margin = Cm(0)

        # 替换这段内容，原本插入图片的部分：
        for img_path in self.image_paths:
            try:
                with Image.open(img_path) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_height / img_width

                    target_width_cm = 21.0  # A4宽度
                    max_height_cm = 29.5  # A4高度

                    target_height_cm = target_width_cm * aspect_ratio

                    # 如果太高，就按高度限制缩小（避免 Word 强制分页）
                    if target_height_cm > max_height_cm:
                        target_height_cm = max_height_cm
                        target_width_cm = target_height_cm / aspect_ratio

                    doc.add_picture(img_path, width=Cm(target_width_cm), height=Cm(target_height_cm))
                    # doc.add_paragraph("")
            except Exception as e:
                wx.MessageBox(f"插入图片失败：{img_path}\n{e}", "错误", wx.ICON_ERROR)

        default_dir = os.path.dirname(self.image_paths[0])
        with wx.FileDialog(self, "保存 Word 文档", defaultDir=default_dir,
                           wildcard="Word 文件 (*.docx)|*.docx",
                           style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
            fileDialog.SetFilename("output.docx")
            if fileDialog.ShowModal() == wx.ID_CANCEL:
                return
            output_path = fileDialog.GetPath()
            doc.save(output_path)
            wx.MessageBox(f"文档已保存到：\n{output_path}", "成功", wx.ICON_INFORMATION)
            
if __name__ == "__main__":
    # 创建应用实例
    app = wx.App(False)
    # 创建主框架窗口
    frame = MainFrame(None)
    # 显示窗口
    frame.Show(True)
    # 启动应用的主循环
    app.MainLoop()